from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "FULL_TEST_SCENARIOS_2026-07-23.md"
OUTPUT = ROOT / "FULL_TEST_SCENARIOS_2026-07-23.docx"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "5B6573"
PAGE_WIDTH = 27.7  # A4 landscape
PAGE_HEIGHT = 19.0
MARGIN = 1.25
CONTENT_WIDTH = PAGE_WIDTH - MARGIN * 2


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(round(width_cm / 2.54 * 1440)))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(round(sum(widths) / 2.54 * 1440)))
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(round(width / 2.54 * 1440)))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side in ("top", "start", "bottom", "end"):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), "80" if side in ("top", "bottom") else "100")
                node.set(qn("w:type"), "dxa")


def set_run_font(run, size, color="000000", bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run_font(run, 8, MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def clean_cell(value):
    value = value.strip()
    value = value.replace("`", "")
    value = value.replace("**", "")
    return value


def widths_for(column_count):
    if column_count == 2:
        return [4.5, CONTENT_WIDTH - 4.5]
    if column_count == 3:
        return [2.7, 4.4, CONTENT_WIDTH - 7.1]
    if column_count == 4:
        return [1.5, 1.7, 5.8, CONTENT_WIDTH - 9.0]
    return [CONTENT_WIDTH / column_count] * column_count


def add_table(doc, rows):
    if not rows or len(rows[0]) < 2:
        return
    count = len(rows[0])
    rows = [row[:count] + [""] * max(0, count - len(row)) for row in rows]
    table = doc.add_table(rows=0, cols=count)
    table.style = "Table Grid"
    for r_index, values in enumerate(rows):
        cells = table.add_row().cells
        for c_index, value in enumerate(values):
            paragraph = cells[c_index].paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(clean_cell(value))
            set_run_font(run, 8.4 if count >= 4 else 8.8, NAVY if r_index == 0 else "000000", bold=r_index == 0)
            if r_index == 0:
                set_cell_shading(cells[c_index], LIGHT_BLUE)
        if r_index == 0:
            set_repeat_table_header(table.rows[-1])
    set_table_geometry(table, widths_for(count))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AUTOWASH PRO")
    set_run_font(run, 12, BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Bộ kịch bản test đầy đủ")
    set_run_font(run, 28, NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("React/Vite + ASP.NET Core .NET 8")
    set_run_font(run, 13, MUTED)
    values = [("Phạm vi", "Manual E2E, API integration và automation backlog"), ("Số lượng", "123 test scenario, 61 case P0"), ("Cập nhật", "23/07/2026")]
    for label, value in values:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(f"{label}: "), 10, NAVY, bold=True)
        set_run_font(p.add_run(value), 10, MUTED)
    doc.add_page_break()


def build():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(PAGE_WIDTH)
    section.page_height = Cm(PAGE_HEIGHT)
    section.top_margin = Cm(MARGIN)
    section.bottom_margin = Cm(MARGIN)
    section.left_margin = Cm(MARGIN)
    section.right_margin = Cm(MARGIN)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    configure_styles(doc)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run("AutoWash Pro | Bộ kịch bản test"), 8, MUTED)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            add_table(doc, table_rows)
            table_rows = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("|") and line.endswith("|"):
            values = [part.strip() for part in line.strip("|").split("|")]
            if all(value and set(value) <= {"-", ":"} for value in values):
                continue
            table_rows.append(values)
            continue
        flush_table()
        if not line or set(line) == {"-"}:
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(clean_cell(line[2:]))
            set_run_font(run, 10)
        elif line[:2].isdigit() and line[1:3] == ". ":
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(clean_cell(line[3:]))
            set_run_font(run, 10)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(clean_cell(line))
            set_run_font(run, 10)
    flush_table()
    doc.core_properties.title = "AutoWash Pro - Bộ kịch bản test đầy đủ"
    doc.core_properties.subject = "Test scenarios"
    doc.core_properties.author = "AutoWash Pro"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
