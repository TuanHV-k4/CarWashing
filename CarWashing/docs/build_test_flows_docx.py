from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
SOURCE = BASE / "FE_BE_DETAILED_TEST_FLOWS.md"
OUTPUT = BASE / "FE_BE_DETAILED_TEST_FLOWS.docx"

CONTENT_WIDTH_DXA = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, col_widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(col_widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in col_widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = col_widths_dxa[min(idx, len(col_widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def split_table_row(line):
    text = line.strip().strip("|")
    return [cell.strip().replace("`", "") for cell in text.split("|")]


def is_separator(line):
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def column_widths(rows):
    count = max(len(r) for r in rows)
    max_lens = [1] * count
    for row in rows:
        for i, cell in enumerate(row):
            max_lens[i] = max(max_lens[i], min(len(cell), 55))
    total = sum(max_lens)
    widths = [max(850, int(CONTENT_WIDTH_DXA * l / total)) for l in max_lens]
    diff = CONTENT_WIDTH_DXA - sum(widths)
    widths[-1] += diff
    return widths


def add_table(doc, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    normalized = [r + [""] * (col_count - len(r)) for r in rows]
    table = doc.add_table(rows=len(normalized), cols=col_count)
    table.style = "Table Grid"
    widths = column_widths(normalized)
    set_table_geometry(table, widths)

    for r_idx, row in enumerate(normalized):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(text)
            run.font.size = Pt(9 if col_count >= 5 else 10)
            if r_idx == 0:
                run.bold = True
                set_cell_shading(cell, "E8EEF5")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif len(text) <= 18 and col_count >= 4:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
    doc.add_paragraph()


def add_paragraph_with_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = p.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9.5)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    setup_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("AutoWash Pro - FE/BE Detailed Test Flows")
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.add_run("Manual end-to-end checklist for React frontend and ASP.NET Core backend mapping.").italic = True
    subtitle.paragraph_format.space_after = Pt(12)

    footer = section.footer.paragraphs[0]
    footer.add_run("AutoWash Pro Test Flows | Page ")
    add_page_number(footer)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        if line.startswith("```"):
            code = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            add_code_block(doc, code)
            i += 1
            continue

        if line.lstrip().startswith("|") and i + 1 < len(lines) and is_separator(lines[i + 1]):
            rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        if line.startswith("# "):
            # Title already added.
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style=None)
            p.style = doc.styles["Normal"]
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            p.paragraph_format.space_after = Pt(4)
            p.add_run("- ")
            p.add_run(line[2:].strip())
            i += 1
            continue

        add_paragraph_with_code(doc, line)
        i += 1

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

