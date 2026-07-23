from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
OUTPUT = BASE / "AUTO_WASH_PRO_TEST_DATA_AND_FLOWS.docx"
CONTENT_WIDTH_DXA = 9360


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            width = widths[min(i, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)


def widths(count):
    return {
        2: [2700, 6660],
        3: [1800, 3000, 4560],
        4: [1300, 2500, 2800, 2760],
        5: [900, 1900, 2300, 2300, 1960],
        6: [750, 1500, 1700, 2100, 1800, 1510],
    }.get(count, [int(CONTENT_WIDTH_DXA / count)] * count)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    geometry(t, widths(len(headers)))
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        shade(c, "E8EEF5")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if len(str(value)) <= 14:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            r.font.size = Pt(9)
    doc.add_paragraph()


def para(doc, text, label=None):
    p = doc.add_paragraph()
    if label:
        r = p.add_run(label)
        r.bold = True
    p.add_run(text)


def bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.add_run("- ")
    p.add_run(text)


def code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(9)
    doc.add_paragraph()


def setup():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.keep_with_next = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("AutoWash Pro - Test data & flows")
    return doc


def flow(doc, title, account, data_rows, steps):
    doc.add_heading(title, level=1)
    para(doc, account, "Tài khoản dùng test: ")
    table(doc, ["Dữ liệu", "Giá trị nên chọn", "Lý do"], data_rows)
    table(doc, ["Bước", "Màn hình", "Thao tác", "API BE", "Kết quả mong đợi"], steps)


def build():
    doc = setup()
    p = doc.add_paragraph()
    r = p.add_run("AutoWash Pro - Data Test Và Luồng Test Chi Tiết")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    para(doc, "Tài liệu này tổng hợp dữ liệu seed và cách test toàn bộ luồng FE-BE mượt nhất, bao gồm account, branch, service, promotion, reward và expected result.")

    doc.add_heading("1. Cấu hình chạy hệ thống", level=1)
    table(doc, ["Thành phần", "Giá trị"], [
        ["Backend đang dùng", "http://localhost:5152"],
        ["Frontend", "http://localhost:5173"],
        ["Swagger", "http://localhost:5152/swagger"],
        ["File seed", "CarWashing-main/database/seed-data.sql"],
        ["File env FE", "CarWashingSystem-UI/.env.local"],
    ])
    code(doc, [
        "cd CarWashing-main/API",
        "dotnet run --launch-profile http",
        "",
        "cd CarWashingSystem-UI",
        "pnpm dev",
        "",
        "# CarWashingSystem-UI/.env.local",
        "VITE_API_BASE_URL=http://localhost:5152",
    ])

    doc.add_heading("2. Account và dữ liệu customer", level=1)
    table(doc, ["Username", "Password", "Role/Tier", "Điểm", "Xe", "Dùng tốt cho"], [
        ["admin", "Admin@123", "Admin", "-", "-", "Tạo/cấu hình dữ liệu và vận hành booking"],
        ["staff", "Staff@123", "Staff", "-", "-", "Vận hành booking"],
        ["demo_customer", "Customer@123", "Silver", "450", "Toyota Vios 51A12345", "Đặt lịch, test promotion cơ bản"],
        ["demo_vip", "Customer@123", "Gold", "2100", "Hyundai Santa Fe 30B67890", "Test reward, bonus points"],
        ["demo_bronze", "Customer@123", "Bronze", "50", "Kia Morning 59C11111", "Test khách ít điểm"],
        ["demo_platinum", "Customer@123", "Platinum", "8500", "Mercedes GLC 51F99999", "Test VIP promotion và reward"],
    ])

    doc.add_heading("3. Catalog seed nên dùng", level=1)
    table(doc, ["Nhóm", "Dữ liệu", "Giá trị test"], [
        ["Branch", "AutoWash Pro - District 1", "Dễ test nhất, có Bay 1 và Bay 2"],
        ["Branch", "AutoWash Pro - Thu Duc", "Dùng thêm với promotion THUDUC20"],
        ["Service", "Basic Wash", "80,000 VND, 20 phút"],
        ["Service", "Premium Wash", "150,000 VND, 35 phút"],
        ["Service", "Interior Clean", "120,000 VND, 30 phút"],
        ["Service", "Full Detail", "350,000 VND, 90 phút"],
    ])

    doc.add_heading("4. Promotion seed đầy đủ", level=1)
    table(doc, ["Mã", "Loại", "Điều kiện", "Service", "Customer nên dùng"], [
        ["JULY15", "Giảm 15%", "Min 100,000", "Premium Wash", "demo_customer"],
        ["GOLDPOINTS", "+200 điểm", "Gold/Platinum, min 300,000", "Full Detail", "demo_vip/demo_platinum"],
        ["WELCOME20K", "Giảm 20,000", "Min 80,000", "Basic/Premium", "demo_customer/demo_bronze"],
        ["BASIC10", "Giảm 10%", "Min 80,000", "Basic Wash", "mọi customer"],
        ["INTERIOR30K", "Giảm 30,000", "Min 120,000", "Interior Clean", "mọi customer"],
        ["SILVER100", "+100 điểm", "Silver trở lên, min 150,000", "Premium/Full Detail", "demo_customer trở lên"],
        ["VIP25", "Giảm 25%", "Platinum, min 300,000", "Full Detail", "demo_platinum"],
        ["FREEBASIC", "Free Basic Wash", "Min 300,000", "Full Detail", "demo_vip/demo_platinum"],
        ["THUDUC20", "Giảm 20%", "Min 100,000", "Premium/Interior", "mọi customer"],
        ["DETAIL100K", "Giảm 100,000", "Min 350,000", "Full Detail", "demo_vip/demo_platinum"],
    ])

    doc.add_heading("5. Reward seed", level=1)
    table(doc, ["Reward", "Loại", "Điểm cần", "Giá trị", "Ai đủ điểm"], [
        ["80K Wash Voucher", "Giảm cố định", "800", "80,000", "demo_vip, demo_platinum"],
        ["Free Basic Wash", "Miễn phí dịch vụ", "1200", "80,000", "demo_vip, demo_platinum"],
    ])

    flow(doc, "6. Luồng Login và Dashboard customer", "demo_customer / Customer@123", [
        ["Customer", "demo_customer", "Có xe, có lịch sử, có điểm"],
        ["Expected tier", "Silver", "Kiểm tra dashboard"],
        ["Expected points", "450", "Kiểm tra loyalty/dashboard"],
    ], [
        ["1", "/login", "Đăng nhập demo_customer", "POST /api/Auth/login", "Redirect /customer/dashboard"],
        ["2", "/customer/dashboard", "Xem KPI", "GET /api/customers/me", "Hiện Silver và 450 điểm"],
        ["3", "/customer/dashboard", "Xem lịch gần nhất", "GET /api/bookings?pageSize=8", "Có dữ liệu nếu có booking"],
        ["4", "/customer/dashboard", "Xem lịch sử rửa", "GET /api/wash-histories/me?pageSize=5", "Thấy booking completed seed"],
    ])

    flow(doc, "7. Luồng đặt lịch không promotion", "demo_customer / Customer@123", [
        ["Xe", "51A12345", "Xe seed sẵn"],
        ["Branch", "AutoWash Pro - District 1", "Có bay hoạt động"],
        ["Service", "Basic Wash", "Giá thấp, dễ test"],
        ["Thời gian", "Thời gian tương lai", "Tránh lỗi slot/ngày quá khứ"],
    ], [
        ["1", "/customer/bookings/new", "Chọn xe/branch/service/time", "GET vehicles/services/branches", "Dropdown có dữ liệu"],
        ["2", "/customer/bookings/new", "Submit", "POST /api/bookings", "Tạo booking Pending"],
        ["3", "/customer/bookings/new", "FE tự tạo payment", "POST /api/payments", "Payment Cash được tạo"],
        ["4", "/customer/bookings/{id}", "Xem chi tiết", "GET /api/bookings/{id}", "Hiện status Pending"],
    ])

    flow(doc, "8. Luồng đặt lịch với promotion dễ nhất", "demo_customer / Customer@123", [
        ["Xe", "51A12345", "Xe seed sẵn"],
        ["Branch", "AutoWash Pro - District 1", "Dễ test"],
        ["Service", "Premium Wash", "Phù hợp JULY15"],
        ["Promotion", "JULY15", "Không yêu cầu tier, min spend đủ"],
    ], [
        ["1", "/customer/bookings/new", "Chọn Premium Wash", "GET /api/services", "Tổng tiền 150,000"],
        ["2", "/customer/bookings/new", "Nhập mã JULY15", "GET /api/loyalty/promotions?pageSize=50", "FE tìm thấy promotion"],
        ["3", "/customer/bookings/new", "Submit", "POST /api/bookings", "Booking Pending"],
        ["4", "/customer/bookings/new", "FE apply mã", "POST /api/loyalty/promotions/{id}/apply", "Áp dụng giảm 15% nếu rule hợp lệ"],
        ["5", "/customer/bookings/{id}", "Kiểm tra detail", "GET /api/bookings/{id}", "Booking vẫn xem được"],
    ])

    flow(doc, "9. Luồng promotion theo tier Gold/Platinum", "demo_vip / Customer@123 hoặc demo_platinum / Customer@123", [
        ["Customer", "demo_vip", "Gold, đủ điều kiện GOLDPOINTS"],
        ["Customer VIP", "demo_platinum", "Platinum, đủ điều kiện VIP25"],
        ["Service", "Full Detail", "Giá 350,000, đủ min spend"],
        ["Promotion", "GOLDPOINTS / DETAIL100K / VIP25", "Test bonus points và discount cao"],
    ], [
        ["1", "/login", "Login demo_vip hoặc demo_platinum", "POST /api/Auth/login", "Vào customer dashboard"],
        ["2", "/customer/bookings/new", "Chọn Full Detail", "GET /api/services", "Tổng 350,000"],
        ["3", "/customer/bookings/new", "Nhập GOLDPOINTS hoặc DETAIL100K", "GET /api/loyalty/promotions", "FE tìm thấy mã"],
        ["4", "/customer/bookings/new", "Submit", "POST /api/bookings + POST /apply", "Promotion được áp dụng nếu đúng tier"],
        ["5", "/customer/loyalty", "Kiểm tra điểm", "GET point balance/history", "Điểm có thể cập nhật theo rule BE"],
    ])

    flow(doc, "10. Luồng Admin vận hành booking", "admin / Admin@123", [
        ["Booking", "Booking Pending vừa tạo", "Cần có booking trước"],
        ["Route", "/admin/bookings", "Board vận hành"],
        ["Thứ tự", "Pending -> Confirmed -> InProgress -> Completed", "Đúng flow nghiệp vụ"],
    ], [
        ["1", "/login", "Login admin", "POST /api/Auth/login", "Redirect /admin/dashboard"],
        ["2", "/admin/bookings", "Tìm booking Pending", "GET /api/bookings?pageSize=100", "Booking nằm ở cột Pending"],
        ["3", "/admin/bookings", "Click Xác nhận", "POST /api/bookings/{id}/confirm", "Status Confirmed"],
        ["4", "/admin/bookings", "Click Bắt đầu", "POST /api/bookings/{id}/start", "Status InProgress"],
        ["5", "/admin/bookings", "Click Hoàn tất", "POST /api/bookings/{id}/complete", "Status Completed"],
    ])

    flow(doc, "11. Luồng xem history sau khi hoàn tất", "demo_customer / Customer@123", [
        ["Điều kiện", "Có booking Completed", "Tạo từ luồng admin"],
        ["Route", "/customer/history", "Customer xem lịch sử"],
        ["Seed có sẵn", "Booking 09/07/2026, 200,000, +85 điểm", "Dùng để đối chiếu"],
    ], [
        ["1", "/customer/history", "Mở lịch sử", "GET /api/wash-histories/me?pageSize=20", "Thấy lịch sử seed và/hoặc booking mới"],
        ["2", "/customer/dashboard", "Kiểm tra mục gần đây", "GET /api/wash-histories/me?pageSize=5", "Hiện lần rửa gần nhất"],
        ["3", "/customer/loyalty", "Kiểm tra điểm", "GET /api/loyalty/customers/{id}/points/balance", "Điểm phản ánh sau complete nếu BE cộng điểm"],
    ])

    flow(doc, "12. Luồng redeem reward", "demo_vip / Customer@123", [
        ["Customer", "demo_vip", "Có 2100 điểm, đủ đổi reward"],
        ["Reward 1", "80K Wash Voucher", "Cần 800 điểm"],
        ["Reward 2", "Free Basic Wash", "Cần 1200 điểm"],
        ["Điều kiện", "Cần booking Pending", "FE yêu cầu chọn booking để redeem"],
    ], [
        ["1", "/customer/bookings/new", "Tạo booking mới nhưng chưa admin confirm", "POST /api/bookings", "Booking Pending"],
        ["2", "/customer/loyalty", "Chọn booking Pending", "GET /api/bookings?status=Pending", "Dropdown có booking"],
        ["3", "/customer/loyalty", "Click đổi 80K Wash Voucher", "POST /api/loyalty/rewards/{id}/redeem", "Redeem thành công nếu đủ điểm"],
        ["4", "/customer/loyalty", "Xem điểm/lịch sử điểm", "GET point balance/history", "Điểm giảm hoặc transaction xuất hiện"],
    ])

    flow(doc, "13. Luồng Admin kiểm tra catalog và tạo thêm data", "admin / Admin@123", [
        ["Route", "/admin/catalog", "Quản lý service/branch/bay"],
        ["Service test", "Express Wash Test - 90,000 - 25 phút", "Tạo thêm nếu cần"],
        ["Branch test", "AutoWash Pro - Test Branch", "Tạo thêm nếu cần"],
    ], [
        ["1", "/admin/catalog", "Mở catalog", "GET services/branches/wash-bays", "Hiện danh sách seed"],
        ["2", "/admin/catalog", "Tạo service mới", "POST /api/services", "Service xuất hiện"],
        ["3", "/admin/catalog", "Tạo branch mới", "POST /api/branches", "Branch xuất hiện"],
        ["4", "/admin/catalog", "Tạo bay mới", "POST /api/wash-bays", "Bay xuất hiện"],
        ["5", "/customer/bookings/new", "Customer kiểm tra dropdown", "GET services/branches", "Data admin tạo xuất hiện"],
    ])

    flow(doc, "14. Luồng Admin logs và dashboard", "admin / Admin@123", [
        ["Dashboard", "/admin/dashboard", "KPI tổng quan"],
        ["Logs", "/admin/logs", "Behavioral logs seed"],
        ["Export", "CSV", "Tải log"],
    ], [
        ["1", "/admin/dashboard", "Mở dashboard", "GET /api/loyalty/dashboard", "KPI hiển thị"],
        ["2", "/admin/dashboard", "Xem chart booking", "GET /api/bookings?pageSize=100", "Chart có dữ liệu"],
        ["3", "/admin/logs", "Mở logs", "GET /api/admin/behavioral-logs?pageSize=50", "Thấy logs seed"],
        ["4", "/admin/logs", "Click Xuất CSV", "GET /api/admin/behavioral-logs/export", "Tải CSV"],
    ])

    flow(doc, "15. Luồng AI customer", "demo_customer / Customer@123", [
        ["Route", "/customer/ai", "AI customer"],
        ["Suggest data", "SUV + nhanh và tiết kiệm", "Gợi ý service"],
        ["Chat data", "Tôi có bao nhiêu điểm?", "Kiểm tra context customer"],
    ], [
        ["1", "/customer/ai", "Chọn loại xe/nhu cầu", "Không gọi API", "Form sẵn sàng"],
        ["2", "/customer/ai", "Click Gợi ý dịch vụ", "POST /api/ai/suggest-services", "Hiện danh sách gợi ý"],
        ["3", "/customer/ai", "Gửi câu hỏi điểm", "POST /api/ai/chat", "AI trả lời"],
        ["4", "/customer/ai", "Gửi prompt injection", "POST /api/ai/chat", "Không lộ API key/prompt"],
    ])

    doc.add_heading("16. Bộ test nhanh khuyến nghị", level=1)
    table(doc, ["Thứ tự", "Luồng", "Data dùng", "Expected"], [
        ["1", "Login customer", "demo_customer", "Vào dashboard"],
        ["2", "Đặt lịch promotion", "Premium Wash + JULY15", "Booking Pending + promotion apply"],
        ["3", "Admin vận hành", "admin", "Booking Completed"],
        ["4", "Customer history", "demo_customer", "Có history mới"],
        ["5", "Reward", "demo_vip + 80K Wash Voucher", "Redeem thành công"],
        ["6", "Promotion VIP", "demo_platinum + Full Detail + VIP25", "Promotion hợp lệ"],
        ["7", "Logs", "admin", "Có behavioral logs"],
        ["8", "AI", "demo_customer", "AI trả lời/gợi ý service"],
    ])

    doc.add_heading("17. Negative cases nên test", level=1)
    table(doc, ["Nhóm", "Data", "Expected"], [
        ["Auth", "Sai mật khẩu", "Không login được"],
        ["Booking", "Bỏ trống xe/service/branch", "FE validation, không gọi POST booking"],
        ["Promotion", "demo_bronze + GOLDPOINTS", "Không đủ tier, apply fail"],
        ["Promotion", "Basic Wash + JULY15", "Sai service, apply fail nếu BE validate service"],
        ["Reward", "demo_customer đổi 80K Voucher", "Không đủ điểm, redeem fail"],
        ["AI", "Không token gọi /api/ai/chat", "401"],
        ["Admin logs", "Customer gọi /api/admin/behavioral-logs", "403"],
    ])

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

