from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("AutoWash_Pro_API_FE_Chi_Tiet.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(242, 244, 247)
LIGHT_BLUE = RGBColor(232, 238, 245)
INK = RGBColor(20, 31, 45)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.bold = bold
    if bold:
        run.font.color.rgb = INK
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbers(doc: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{idx}.  ")
        r.bold = True
        p.add_run(item)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    for line_idx, line in enumerate(text.strip("\n").splitlines()):
        if line_idx:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8.5)


def add_endpoint_table(doc: Document, rows: list[tuple[str, str, str, str]], widths: list[float] | None = None) -> None:
    widths = widths or [0.85, 2.15, 1.3, 2.2]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Method", "Endpoint", "Auth", "FE dùng khi"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for method, endpoint, auth, purpose in rows:
        cells = table.add_row().cells
        for i, value in enumerate([method, endpoint, auth, purpose]):
            set_cell_text(cells[i], value)
    set_table_widths(table, widths)
    doc.add_paragraph()


def add_fields_table(doc: Document, title: str, rows: list[tuple[str, str, str]]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, header in enumerate(["Field", "Kiểu", "Ghi chú FE"]):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for field, typ, note in rows:
        cells = table.add_row().cells
        for i, value in enumerate([field, typ, note]):
            set_cell_text(cells[i], value)
    set_table_widths(table, [1.75, 1.25, 3.5])
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F4F6F9")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p.add_run("\n" + body)
    set_table_widths(table, [6.5])
    doc.add_paragraph()


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    return doc


def build() -> None:
    doc = setup_doc()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("AutoWash Pro - API Contract va Luong Tich Hop FE")
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(11, 37, 69)
    sub = doc.add_paragraph()
    sub.add_run("Tai lieu danh cho Frontend: endpoint, payload, response, auth, enum, luong man hinh va cac luu y tich hop.").italic = True
    add_callout(
        doc,
        "Pham vi va nguon doi chieu",
        "Tai lieu duoc tong hop tu controller, DTO, enum va service logic trong repo hien tai. Base URL local thuong la https://localhost:<port> hoac http://localhost:<port>. Swagger co tai /swagger khi API chay.",
    )

    add_heading(doc, "1. Quy uoc chung", 1)
    add_bullets(
        doc,
        [
            "Request/response mac dinh la JSON. Header chung: Content-Type: application/json.",
            "Endpoint co auth can them Authorization: Bearer <accessToken>.",
            "Enum serialize theo ten chuoi do backend dung JsonStringEnumConverter; FE gui chuoi enum, khong gui so.",
            "DateTime gui theo ISO 8601, uu tien UTC, vi du 2026-07-01T09:00:00Z. TimeSpan cua branch gui dang 08:00:00.",
            "Loi validation/business cua Operations/Loyalty thuong tra application/problem+json voi title/detail/status.",
            "Operations va Loyalty hien khong gan [Authorize] tren controller, nhung mot so ham booking van doc current customer tu JWT neu co. FE nen van gui token sau login cho cac luong customer.",
        ],
    )
    add_code(
        doc,
        """Authorization: Bearer <accessToken>
Content-Type: application/json

ProblemDetails mau:
{
  "type": "https://tools.ietf.org/html/rfc9110#section-15.5.1",
  "title": "Validation Error",
  "status": 400,
  "detail": "Booking time must be in the future."
}""",
    )

    add_heading(doc, "2. Auth va phan quyen", 1)
    add_endpoint_table(
        doc,
        [
            ("POST", "/api/Auth/register", "Public", "Dang ky customer, backend tao user/customer va gui OTP email."),
            ("POST", "/api/Auth/verify-email", "Public", "Xac thuc email bang OTP."),
            ("POST", "/api/Auth/resend-otp", "Public", "Gui lai OTP khi het han/khong nhan duoc."),
            ("POST", "/api/Auth/login", "Public", "Dang nhap, nhan JWT de FE luu va gan vao header."),
            ("GET", "/api/Auth/me", "Bearer", "Kiem tra token va lay claims user hien tai."),
            ("GET", "/api/Auth/admin-only", "AdminOnly", "Test role Admin."),
            ("GET", "/api/Auth/customer-only", "CustomerOnly", "Test role Customer."),
        ],
    )
    add_fields_table(
        doc,
        "Auth DTO chinh",
        [
            ("RegisterRequest", "object", "username, password, confirmPassword, fullName, email, phoneNumber?"),
            ("LoginRequest", "object", "username, password"),
            ("LoginResponse", "object", "userID, username, fullName, email, role, accessToken, accessTokenExpiration"),
            ("VerifyEmailRequest", "object", "email, otpCode"),
        ],
    )
    add_code(
        doc,
        """POST /api/Auth/login
{
  "username": "customer01",
  "password": "Password123!"
}

Response:
{
  "userID": "guid",
  "username": "customer01",
  "fullName": "Nguyen Van A",
  "email": "a@example.com",
  "role": "Customer",
  "accessToken": "jwt",
  "accessTokenExpiration": "2026-07-01T00:00:00Z"
}""",
    )
    add_heading(doc, "Luong FE de xac thuc", 2)
    add_numbers(
        doc,
        [
            "Man hinh Register goi POST /api/Auth/register; neu thanh cong chuyen sang man hinh OTP.",
            "Man hinh OTP goi POST /api/Auth/verify-email. Neu OTP sai/hết hạn hien message detail; co nut goi resend-otp.",
            "Man hinh Login goi POST /api/Auth/login, luu accessToken va role. Moi request can auth gan Authorization header.",
            "Khi reload app, FE co the goi GET /api/Auth/me de verify token va khoi phuc session.",
        ],
    )

    add_heading(doc, "3. Customer va Vehicle", 1)
    add_endpoint_table(
        doc,
        [
            ("GET", "/api/customers/me", "CustomerOnly", "Lay profile, diem, tier hien tai va tier perks."),
            ("GET", "/api/vehicles/me", "CustomerOnly", "Danh sach xe cua customer hien tai."),
            ("POST", "/api/vehicles", "CustomerOnly", "Them xe."),
            ("PUT", "/api/vehicles/{vehicleId}", "CustomerOnly", "Sua thong tin xe."),
            ("PUT", "/api/vehicles/{vehicleId}/status", "CustomerOnly", "Doi status xe."),
        ],
    )
    add_fields_table(
        doc,
        "Vehicle DTO",
        [
            ("licensePlate", "string", "Bien so, FE nen normalize upper-case neu can."),
            ("vehicleType", "VehicleTypeEnum?", "Sedan, SUV, Hatchback, Coupe, Convertible, PickupTruck, Van, Motorcycle."),
            ("brand/model/color", "string?", "Thong tin mo ta xe."),
            ("status", "VehicleStatusEnum", "Active, Inactive, Suspended."),
        ],
    )
    add_code(
        doc,
        """POST /api/vehicles
{
  "licensePlate": "51A12345",
  "vehicleType": "Sedan",
  "brand": "Toyota",
  "model": "Vios",
  "color": "White"
}""",
    )

    add_heading(doc, "4. Operations: catalog, chi nhanh, bay rua", 1)
    add_endpoint_table(
        doc,
        [
            ("GET", "/api/services?page=1&pageSize=20&includeInactive=false", "None/Bearer", "Lay goi dich vu dang ban."),
            ("GET", "/api/services/{id}", "None/Bearer", "Chi tiet dich vu."),
            ("POST", "/api/services", "None/Bearer", "Tao dich vu."),
            ("PUT", "/api/services/{id}", "None/Bearer", "Sua dich vu."),
            ("DELETE", "/api/services/{id}", "None/Bearer", "Inactive/archive dich vu."),
            ("GET", "/api/branches?page=1&pageSize=20&includeInactive=false", "None/Bearer", "Lay chi nhanh dang mo."),
            ("GET", "/api/wash-bays?page=1&pageSize=20&branchId={guid}", "None/Bearer", "Lay bay rua theo chi nhanh."),
        ],
    )
    add_fields_table(
        doc,
        "Service / Branch / WashBay request",
        [
            ("CreateServiceRequest", "object", "name, description?, price, durationMinutes"),
            ("UpdateServiceRequest", "object", "name, description?, price, durationMinutes, isActive"),
            ("CreateBranchRequest", "object", "name, address, phone?, openTime, closeTime"),
            ("CreateWashBayRequest", "object", "branchId, name"),
        ],
    )
    add_callout(
        doc,
        "Canh bao route trung api/services",
        "Repo hien co them controller legacy API/Controllers/ServicesController.cs voi route api/[controller] => /api/Services. ASP.NET route thuong khong phan biet hoa/thuong, co the trung voi /api/services cua Operations. FE nen uu tien Operations Services va backend nen doi/xoa route legacy de tranh ambiguous endpoint.",
    )

    doc.add_page_break()
    add_heading(doc, "5. Bookings", 1)
    add_endpoint_table(
        doc,
        [
            ("GET", "/api/bookings?status=Pending&fromDate=...&toDate=...&branchId=...&page=1&pageSize=20", "None/Bearer", "Danh sach booking; customer chi thay cua minh neu token co customerId, Admin/Staff thay rong hon."),
            ("GET", "/api/bookings/{id}", "None/Bearer", "Chi tiet booking kem branchName/washBayName."),
            ("POST", "/api/bookings", "Bearer nen co", "Customer tao lich rua xe."),
            ("POST", "/api/bookings/{id}/cancel", "Bearer nen co", "Huy booking khi chua completed va chua paid."),
            ("POST", "/api/bookings/{id}/confirm", "Staff/Admin flow", "Pending -> Confirmed."),
            ("POST", "/api/bookings/{id}/start", "Staff/Admin flow", "Confirmed -> InProgress."),
            ("POST", "/api/bookings/{id}/complete", "Staff/Admin flow", "InProgress -> Completed, tao wash history va cong diem."),
        ],
    )
    add_fields_table(
        doc,
        "CreateBookingRequest / BookingResponse",
        [
            ("vehicleId", "guid", "Phai thuoc customer hien tai, lay tu GET /api/vehicles/me."),
            ("branchId", "guid", "Chi nhanh phai Open, lay tu GET /api/branches."),
            ("serviceId", "guid", "Service phai Active, lay tu GET /api/services."),
            ("bookingStartTime", "DateTime", "Phai o tuong lai, nam trong gio mo cua branch, khong vuot bookingWindowDays theo tier."),
            ("washBayId", "guid?", "Backend tu tim bay rua kha dung; FE khong can chon neu luong hien tai khong yeu cau."),
            ("status", "BookingStatusEnum", "Pending, Confirmed, InProgress, Completed, Cancelled, NoShow."),
        ],
    )
    add_code(
        doc,
        """POST /api/bookings
{
  "vehicleId": "guid",
  "branchId": "guid",
  "serviceId": "guid",
  "bookingStartTime": "2026-07-01T09:00:00Z",
  "note": "optional"
}""",
    )
    add_heading(doc, "Luong dat lich FE", 2)
    add_numbers(
        doc,
        [
            "Load profile/customer va danh sach xe: GET /api/customers/me, GET /api/vehicles/me.",
            "Load catalog: GET /api/services va GET /api/branches; khi chon branch co the load GET /api/wash-bays de hien tinh trang noi bo neu can.",
            "FE tinh preview end time = bookingStartTime + durationMinutes cua service, nhung backend moi la nguon chinh xac.",
            "Goi POST /api/bookings. Neu 409 'No available wash bay...' hien thong bao chon khung gio khac.",
            "Sau khi tao booking, co the hien trang thanh toan hoac trang chi tiet booking.",
        ],
    )
    add_heading(doc, "Luong trang thai booking", 2)
    add_code(doc, "Pending -> Confirmed -> InProgress -> Completed\nPending/Confirmed/InProgress -> Cancelled neu chua paid va chua completed")

    add_heading(doc, "6. Payments", 1)
    add_endpoint_table(
        doc,
        [
            ("GET", "/api/payments/{id}", "None/Bearer", "Lay chi tiet payment."),
            ("POST", "/api/payments", "None/Bearer", "Tao payment pending cho booking."),
            ("POST", "/api/payments/{id}/paid", "Staff/Admin flow", "Danh dau da thanh toan."),
            ("POST", "/api/payments/{id}/void", "Staff/Admin flow", "Huy payment pending."),
        ],
    )
    add_fields_table(
        doc,
        "Payment rules",
        [
            ("method", "PaymentMethodEnum", "Cash, CardAtCounter, BankTransfer, EWalletAtCounter. Code error message cu co the ghi Card nhung enum thuc te khong co Card."),
            ("amount", "decimal", "Phai > 0 va bang booking.EstimatedTotalAmount."),
            ("status", "PaymentStatusEnum", "Pending, Paid, Failed, Voided."),
            ("paid", "transition", "Chi Pending moi mark paid; Voided khong paid lai duoc."),
            ("void", "transition", "Chi Pending moi void; Paid khong void duoc."),
        ],
    )
    add_code(
        doc,
        """POST /api/payments
{
  "bookingId": "guid",
  "amount": 80000,
  "method": "Cash",
  "note": "Pay at counter"
}""",
    )

    add_heading(doc, "7. Loyalty, rewards va promotions", 1)
    add_endpoint_table(
        doc,
        [
            ("GET", "/api/loyalty/settings", "None/Bearer", "Lay rule tinh diem."),
            ("GET", "/api/loyalty/tiers", "None/Bearer", "Danh sach tier."),
            ("POST/PUT/DELETE", "/api/loyalty/tiers/{id?}", "Admin flow", "Quan tri tier."),
            ("GET", "/api/loyalty/customers/{customerId}/points/balance", "None/Bearer", "So du diem customer."),
            ("GET", "/api/loyalty/customers/{customerId}/points/history", "None/Bearer", "Lich su diem."),
            ("GET", "/api/loyalty/rewards", "None/Bearer", "Danh sach reward doi diem."),
            ("POST", "/api/loyalty/rewards/{id}/redeem", "Customer/Admin flow", "Doi diem lay reward."),
            ("POST", "/api/loyalty/customers/{customerId}/tier/evaluate", "Admin flow", "Danh gia lai tier."),
            ("GET", "/api/loyalty/dashboard", "Admin flow", "Chi so loyalty."),
            ("GET/POST/PUT/DELETE", "/api/loyalty/promotions", "Admin flow", "CRUD promotion."),
            ("POST", "/api/loyalty/promotions/{id}/send", "Admin flow", "Gan promotion cho danh sach customer."),
            ("POST", "/api/loyalty/promotions/{id}/apply", "Customer/Admin flow", "Ap promotion vao booking pending."),
        ],
    )
    add_fields_table(
        doc,
        "Loyalty rules can biet",
        [
            ("Earn points", "rule", "1 point / 10,000 VND * pointMultiplier cua tier + bonusPoints tu promotion."),
            ("Expiry", "rule", "Diem earn het han sau 12 thang."),
            ("Redeem", "rule", "Tieu diem theo cac transaction Earn con remainingPoints, sap xep theo expiry som nhat."),
            ("Reward type", "enum", "FixedDiscount, PercentageDiscount, FreeService, AddOnService."),
            ("Promotion type", "enum", "PercentageDiscount, FixedDiscount, FreeService, BonusPoints."),
            ("Tier qualification", "enum", "AllConditions hoac AnyCondition."),
        ],
    )
    add_code(
        doc,
        """POST /api/loyalty/promotions/{id}/apply
{
  "bookingId": "guid",
  "customerId": "guid",
  "code": "SUMMER10"
}

Response:
{
  "bookingId": "guid",
  "promotionId": "guid",
  "discountAmount": 50000,
  "bonusPoints": 0,
  "totalBeforeDiscount": 150000,
  "totalAfterDiscount": 100000
}""",
    )
    add_heading(doc, "Luong loyalty khi hoan tat rua xe", 2)
    add_numbers(
        doc,
        [
            "Staff/Admin confirm va start booking theo workflow operations.",
            "Khi complete booking, backend tao WashHistory, tinh discount/bonus tu BookingPromotions, cong diem cho customer.",
            "Backend ghi LoyaltyPointTransaction Earn voi idempotencyKey wash:{bookingId}:earn de tranh cong lap.",
            "Backend cap nhat TotalSpent, TotalVisits, LastVisitDate va evaluate tier.",
            "FE refresh GET /api/customers/me, points/balance, wash-history de cap nhat UI.",
        ],
    )

    add_heading(doc, "8. Wash histories", 1)
    add_endpoint_table(
        doc,
        [
            ("GET", "/api/wash-histories/me?page=1&pageSize=10", "CustomerOnly", "Lich su rua xe cua customer dang dang nhap."),
            ("GET", "/api/wash-histories/me/{washHistoryId}", "CustomerOnly", "Chi tiet mot wash history cua customer."),
            ("GET", "/api/wash-histories/customer/{customerId}?page=1&pageSize=10", "StaffOrAdmin", "Staff/Admin xem lich su theo customer."),
            ("GET", "/api/loyalty/wash-history?customerId={guid}&page=1&pageSize=20", "None/Bearer", "History trong module loyalty; neu khong truyen customerId tra tat ca."),
        ],
    )
    add_fields_table(
        doc,
        "WashHistory fields",
        [
            ("washHistoryID/id", "guid", "Ten field khac nhau giua controller wash-histories va loyalty/wash-history."),
            ("actualTotalAmount", "decimal", "Tong truoc discount trong DTO loyalty."),
            ("discountAmount", "decimal", "Tong discount tu promotion/reward."),
            ("finalAmount", "decimal", "So tien tinh diem va hien thanh toan cuoi."),
            ("pointsEarned", "int", "Diem nhan duoc khi complete."),
        ],
    )

    add_heading(doc, "9. AI", 1)
    add_endpoint_table(
        doc,
        [
            ("POST", "/api/ai/chat", "CustomerOnly + rate limit", "Chat tu van cho customer."),
            ("POST", "/api/ai/suggest-services", "CustomerOnly + rate limit", "Goi y goi dich vu theo xe/nhu cau."),
            ("POST", "/api/ai/admin/chat", "AdminOnly + rate limit", "Chat insight admin."),
        ],
    )
    add_code(
        doc,
        """POST /api/ai/suggest-services
{
  "vehicleType": "Sedan",
  "preference": "cheap and fast"
}

Response:
{
  "suggestions": [
    { "serviceId": "guid", "serviceName": "Basic Wash", "price": 80000, "reason": "..." }
  ],
  "summary": "...",
  "isFallback": false,
  "source": "gemini"
}""",
    )

    add_heading(doc, "10. Admin users va behavioral logs", 1)
    add_endpoint_table(
        doc,
        [
            ("PUT", "/api/admin/users/{userId}/status", "AdminOnly", "Doi trang thai user."),
            ("GET", "/api/admin/behavioral-logs?customerID=...&actionType=Book&from=...&to=...&page=1&pageSize=20", "AdminOnly", "Xem log hanh vi."),
            ("GET", "/api/admin/behavioral-logs/export?... ", "AdminOnly", "Tai CSV behavioral-logs.csv."),
        ],
    )
    add_fields_table(
        doc,
        "Admin enum/filter",
        [
            ("UserStatusEnum", "enum", "Active, Inactive, Suspended, Deleted."),
            ("BehavioralActionTypeEnum", "enum", "ViewPromotion, Book, CancelBooking, LeaveFeedback, RedeemReward."),
            ("BehavioralLogFilterDto", "object", "customerID?, actionType?, from?, to?, page=1, pageSize=20."),
        ],
    )

    add_heading(doc, "11. Quick reference endpoint matrix", 1)
    add_endpoint_table(
        doc,
        [
            ("Auth", "/api/Auth/*", "Mixed", "Dang ky, OTP, login, me."),
            ("Customer", "/api/customers/me", "CustomerOnly", "Profile/tier/points tong quan."),
            ("Vehicles", "/api/vehicles/*", "CustomerOnly", "CRUD xe customer."),
            ("Operations", "/api/services, /api/branches, /api/wash-bays", "None/Bearer", "Catalog/branch/bay."),
            ("Bookings", "/api/bookings/*", "Bearer nen co", "Dat lich va van hanh trang thai."),
            ("Payments", "/api/payments/*", "None/Bearer", "Thu tien tai quay/void."),
            ("Loyalty", "/api/loyalty/*", "None/Bearer", "Tier, diem, reward, promotion."),
            ("History", "/api/wash-histories/*", "Customer/Staff/Admin", "Lich su rua xe."),
            ("AI", "/api/ai/*", "Customer/Admin", "Chat/goi y."),
            ("Admin", "/api/admin/*", "AdminOnly", "Quan tri user/log."),
        ],
        widths=[1.1, 2.6, 1.2, 1.6],
    )

    add_heading(doc, "12. Phu luc enum FE nen map", 1)
    add_fields_table(
        doc,
        "Enums",
        [
            ("VehicleTypeEnum", "string", "Sedan, SUV, Hatchback, Coupe, Convertible, PickupTruck, Van, Motorcycle"),
            ("VehicleStatusEnum", "string", "Active, Inactive, Suspended"),
            ("BookingStatusEnum", "string", "Pending, Confirmed, InProgress, Completed, Cancelled, NoShow"),
            ("PaymentMethodEnum", "string", "Cash, CardAtCounter, BankTransfer, EWalletAtCounter"),
            ("PaymentStatusEnum", "string", "Pending, Paid, Failed, Voided"),
            ("ServiceStatusEnum", "string", "Active, Inactive, Archived"),
            ("BranchStatusEnum", "string", "Open, Closed, UnderMaintenance"),
            ("WashBayStatusEnum", "string", "Active, Inactive, UnderMaintenance"),
            ("LoyaltyTierStatusEnum", "string", "Active, Inactive"),
            ("RewardTypeEnum", "string", "FixedDiscount, PercentageDiscount, FreeService, AddOnService"),
            ("PromotionTypeEnum", "string", "PercentageDiscount, FixedDiscount, FreeService, BonusPoints"),
            ("PromotionStatusEnum", "string", "Draft, Active, Expired, Disabled"),
            ("TierQualificationModeEnum", "string", "AllConditions, AnyCondition"),
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("AutoWash Pro API FE Guide")

    doc.save(OUT)


if __name__ == "__main__":
    build()
