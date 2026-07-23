from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DOCX_PATH = Path("AutoWash_Pro_API_FE_Chi_Tiet.docx")
APPENDIX_TXT_PATH = Path("artifacts/api_json_appendix.txt")


APPENDIX = r"""
13. Phụ lục JSON chi tiết cho từng API

Quy ước chung

Header cho API cần đăng nhập:
```http
Authorization: Bearer <accessToken>
Content-Type: application/json
```

Response lỗi nghiệp vụ/validation thường dùng ProblemDetails:
```json
{
  "type": "https://tools.ietf.org/html/rfc9110#section-15.5.1",
  "title": "Validation Error",
  "status": 400,
  "detail": "Nội dung lỗi cụ thể",
  "instance": "/api/example"
}
```

Response wrapper `ApiResponse<T>`:
```json
{
  "success": true,
  "data": {},
  "message": null
}
```

Response phân trang `PagedResult<T>`:
```json
{
  "items": [],
  "page": 1,
  "pageSize": 20,
  "totalCount": 0,
  "totalPages": 0,
  "hasNext": false,
  "hasPrevious": false
}
```

13.1 Auth

POST /api/Auth/register
Auth: Public
Request:
```json
{
  "username": "customer01",
  "password": "Password123!",
  "confirmPassword": "Password123!",
  "fullName": "Nguyen Van A",
  "email": "a@example.com",
  "phoneNumber": "0900000000"
}
```
Response 201:
```json
{
  "userID": "guid",
  "username": "customer01",
  "fullName": "Nguyen Van A",
  "email": "a@example.com",
  "phoneNumber": "0900000000",
  "role": "Customer",
  "createdAt": "2026-07-01T00:00:00Z"
}
```

POST /api/Auth/verify-email
Auth: Public
Request:
```json
{
  "email": "a@example.com",
  "otpCode": "123456"
}
```
Response 200:
```json
{
  "success": true,
  "data": true,
  "message": "Email verified successfully"
}
```

POST /api/Auth/resend-otp
Auth: Public
Request:
```json
{
  "email": "a@example.com"
}
```
Response 200:
```json
{
  "success": true,
  "data": true,
  "message": "OTP resent successfully"
}
```

POST /api/Auth/login
Auth: Public
Request:
```json
{
  "username": "customer01",
  "password": "Password123!"
}
```
Response 200:
```json
{
  "userID": "guid",
  "username": "customer01",
  "fullName": "Nguyen Van A",
  "email": "a@example.com",
  "role": "Customer",
  "accessToken": "jwt",
  "accessTokenExpiration": "2026-07-01T00:00:00Z"
}
```

GET /api/Auth/me
Auth: Bearer
Request body: không có
Response 200:
```json
{
  "userID": "guid",
  "username": "customer01",
  "email": "a@example.com",
  "fullName": "Nguyen Van A",
  "role": "Customer"
}
```

GET /api/Auth/admin-only
Auth: AdminOnly
Request body: không có
Response 200:
```json
{
  "message": "Admin access granted"
}
```

GET /api/Auth/customer-only
Auth: CustomerOnly
Request body: không có
Response 200:
```json
{
  "message": "Customer access granted"
}
```

13.2 Customer

GET /api/customers/me
Auth: CustomerOnly
Request body: không có
Response 200:
```json
{
  "success": true,
  "data": {
    "customerID": "guid",
    "userID": "guid",
    "username": "customer01",
    "fullName": "Nguyen Van A",
    "email": "a@example.com",
    "phoneNumber": "0900000000",
    "currentPoints": 100,
    "lifetimePoints": 200,
    "totalSpent": 500000,
    "totalVisits": 5,
    "lastVisitDate": "2026-07-01T00:00:00Z",
    "tierID": "guid",
    "tierName": "Silver",
    "tierRank": 2,
    "currentTierSince": "2026-07-01T00:00:00Z",
    "createdAt": "2026-07-01T00:00:00Z",
    "tierPerks": ["5% discount"]
  },
  "message": null
}
```

13.3 Vehicles

GET /api/vehicles/me
Auth: CustomerOnly
Request body: không có
Response 200:
```json
{
  "success": true,
  "data": [
    {
      "vehicleID": "guid",
      "customerID": "guid",
      "licensePlate": "51A12345",
      "vehicleType": "Sedan",
      "brand": "Toyota",
      "model": "Vios",
      "color": "White",
      "status": "Active",
      "createdAt": "2026-07-01T00:00:00Z"
    }
  ],
  "message": null
}
```

POST /api/vehicles
Auth: CustomerOnly
Request:
```json
{
  "licensePlate": "51A12345",
  "vehicleType": "Sedan",
  "brand": "Toyota",
  "model": "Vios",
  "color": "White"
}
```
Response 201:
```json
{
  "success": true,
  "data": {
    "vehicleID": "guid",
    "customerID": "guid",
    "licensePlate": "51A12345",
    "vehicleType": "Sedan",
    "brand": "Toyota",
    "model": "Vios",
    "color": "White",
    "status": "Active",
    "createdAt": "2026-07-01T00:00:00Z"
  },
  "message": "Vehicle created successfully"
}
```

PUT /api/vehicles/{vehicleId}
Auth: CustomerOnly
Request:
```json
{
  "vehicleType": "SUV",
  "brand": "Honda",
  "model": "CR-V",
  "color": "Black"
}
```
Response 200:
```json
{
  "success": true,
  "data": {
    "vehicleID": "guid",
    "customerID": "guid",
    "licensePlate": "51A12345",
    "vehicleType": "SUV",
    "brand": "Honda",
    "model": "CR-V",
    "color": "Black",
    "status": "Active",
    "createdAt": "2026-07-01T00:00:00Z"
  },
  "message": "Vehicle updated successfully"
}
```

PUT /api/vehicles/{vehicleId}/status
Auth: CustomerOnly
Request:
```json
{
  "status": "Inactive"
}
```
Response 200:
```json
{
  "success": true,
  "data": {
    "vehicleID": "guid",
    "customerID": "guid",
    "licensePlate": "51A12345",
    "vehicleType": "SUV",
    "brand": "Honda",
    "model": "CR-V",
    "color": "Black",
    "status": "Inactive",
    "createdAt": "2026-07-01T00:00:00Z"
  },
  "message": "Vehicle status updated successfully"
}
```

13.4 Operations - Services

GET /api/services?page=1&pageSize=20&includeInactive=false
Auth: None/Bearer
Request body: không có
Response 200:
```json
{
  "items": [
    {
      "id": "guid",
      "name": "Basic Wash",
      "description": "Exterior wash",
      "price": 80000,
      "durationMinutes": 30,
      "isActive": true
    }
  ],
  "page": 1,
  "pageSize": 20,
  "totalCount": 1
}
```

GET /api/services/{id}
Auth: None/Bearer
Request body: không có
Response 200:
```json
{
  "id": "guid",
  "name": "Basic Wash",
  "description": "Exterior wash",
  "price": 80000,
  "durationMinutes": 30,
  "isActive": true
}
```

POST /api/services
Auth: None/Bearer
Request:
```json
{
  "name": "Premium Wash",
  "description": "Exterior and interior wash",
  "price": 150000,
  "durationMinutes": 60
}
```
Response 201:
```json
{
  "id": "guid",
  "name": "Premium Wash",
  "description": "Exterior and interior wash",
  "price": 150000,
  "durationMinutes": 60,
  "isActive": true
}
```

PUT /api/services/{id}
Auth: None/Bearer
Request:
```json
{
  "name": "Premium Wash",
  "description": "Exterior, interior and wax",
  "price": 180000,
  "durationMinutes": 75,
  "isActive": true
}
```
Response 200:
```json
{
  "id": "guid",
  "name": "Premium Wash",
  "description": "Exterior, interior and wax",
  "price": 180000,
  "durationMinutes": 75,
  "isActive": true
}
```

DELETE /api/services/{id}
Auth: None/Bearer
Request body: không có
Response 200:
```json
true
```

13.5 Operations - Branches

GET /api/branches?page=1&pageSize=20&includeInactive=false
Auth: None/Bearer
Request body: không có
Response 200:
```json
{
  "items": [
    {
      "id": "guid",
      "name": "District 1",
      "address": "123 Street",
      "phone": "0900000000",
      "openTime": "08:00:00",
      "closeTime": "20:00:00",
      "isActive": true
    }
  ],
  "page": 1,
  "pageSize": 20,
  "totalCount": 1
}
```

GET /api/branches/{id}
Auth: None/Bearer
Request body: không có
Response 200:
```json
{
  "id": "guid",
  "name": "District 1",
  "address": "123 Street",
  "phone": "0900000000",
  "openTime": "08:00:00",
  "closeTime": "20:00:00",
  "isActive": true
}
```

POST /api/branches
Auth: None/Bearer
Request:
```json
{
  "name": "District 1",
  "address": "123 Street",
  "phone": "0900000000",
  "openTime": "08:00:00",
  "closeTime": "20:00:00"
}
```
Response 201:
```json
{
  "id": "guid",
  "name": "District 1",
  "address": "123 Street",
  "phone": "0900000000",
  "openTime": "08:00:00",
  "closeTime": "20:00:00",
  "isActive": true
}
```

PUT /api/branches/{id}
Auth: None/Bearer
Request:
```json
{
  "name": "District 1",
  "address": "456 Street",
  "phone": "0911111111",
  "openTime": "08:00:00",
  "closeTime": "21:00:00",
  "isActive": true
}
```
Response 200:
```json
{
  "id": "guid",
  "name": "District 1",
  "address": "456 Street",
  "phone": "0911111111",
  "openTime": "08:00:00",
  "closeTime": "21:00:00",
  "isActive": true
}
```

DELETE /api/branches/{id}
Auth: None/Bearer
Request body: không có
Response 200:
```json
true
```

13.6 Operations - Wash Bays

GET /api/wash-bays?page=1&pageSize=20&branchId={guid}&includeInactive=false
Auth: None/Bearer
Request body: không có
Response 200:
```json
{
  "items": [
    {
      "id": "guid",
      "branchId": "guid",
      "name": "Bay 1",
      "isActive": true
    }
  ],
  "page": 1,
  "pageSize": 20,
  "totalCount": 1
}
```

GET /api/wash-bays/{id}
Auth: None/Bearer
Request body: không có
Response 200:
```json
{
  "id": "guid",
  "branchId": "guid",
  "name": "Bay 1",
  "isActive": true
}
```

POST /api/wash-bays
Auth: None/Bearer
Request:
```json
{
  "branchId": "guid",
  "name": "Bay 1"
}
```
Response 201:
```json
{
  "id": "guid",
  "branchId": "guid",
  "name": "Bay 1",
  "isActive": true
}
```

PUT /api/wash-bays/{id}
Auth: None/Bearer
Request:
```json
{
  "branchId": "guid",
  "name": "Bay 1 Updated",
  "isActive": true
}
```
Response 200:
```json
{
  "id": "guid",
  "branchId": "guid",
  "name": "Bay 1 Updated",
  "isActive": true
}
```

DELETE /api/wash-bays/{id}
Auth: None/Bearer
Request body: không có
Response 200:
```json
true
```

13.7 Bookings

GET /api/bookings?status=Pending&fromDate=2026-07-01&toDate=2026-07-31&branchId={guid}&page=1&pageSize=20
Auth: None/Bearer; FE nên gửi Bearer cho luồng customer
Request body: không có
Response 200:
```json
{
  "items": [
    {
      "id": "guid",
      "customerId": "guid",
      "vehicleId": "guid",
      "branchId": "guid",
      "serviceId": "guid",
      "washBayId": "guid",
      "bookingStartTime": "2026-07-01T09:00:00Z",
      "bookingEndTime": "2026-07-01T09:30:00Z",
      "status": "Pending",
      "totalAmount": 80000,
      "serviceNameSnapshot": "Basic Wash",
      "durationMinutesSnapshot": 30,
      "priceSnapshot": 80000,
      "tierSnapshot": "Silver",
      "createdAt": "2026-07-01T00:00:00Z",
      "note": "optional"
    }
  ],
  "page": 1,
  "pageSize": 20,
  "totalCount": 1
}
```

GET /api/bookings/{id}
Auth: None/Bearer
Request body: không có
Response 200:
```json
{
  "id": "guid",
  "customerId": "guid",
  "vehicleId": "guid",
  "branchId": "guid",
  "serviceId": "guid",
  "washBayId": "guid",
  "bookingStartTime": "2026-07-01T09:00:00Z",
  "bookingEndTime": "2026-07-01T09:30:00Z",
  "status": "Pending",
  "totalAmount": 80000,
  "serviceNameSnapshot": "Basic Wash",
  "durationMinutesSnapshot": 30,
  "priceSnapshot": 80000,
  "tierSnapshot": "Silver",
  "createdAt": "2026-07-01T00:00:00Z",
  "note": "optional",
  "branchName": "District 1",
  "washBayName": "Bay 1"
}
```

POST /api/bookings
Auth: Bearer nên có customerId claim
Request:
```json
{
  "vehicleId": "guid",
  "branchId": "guid",
  "serviceId": "guid",
  "bookingStartTime": "2026-07-01T09:00:00Z",
  "note": "optional"
}
```
Response 201:
```json
{
  "id": "guid",
  "customerId": "guid",
  "vehicleId": "guid",
  "branchId": "guid",
  "serviceId": "guid",
  "washBayId": "guid",
  "bookingStartTime": "2026-07-01T09:00:00Z",
  "bookingEndTime": "2026-07-01T09:30:00Z",
  "status": "Pending",
  "totalAmount": 80000,
  "serviceNameSnapshot": "Basic Wash",
  "durationMinutesSnapshot": 30,
  "priceSnapshot": 80000,
  "tierSnapshot": "Silver",
  "createdAt": "2026-07-01T00:00:00Z",
  "note": "optional"
}
```

POST /api/bookings/{id}/cancel
Auth: Bearer nên có; Admin/Staff có thể hủy rộng hơn
Request:
```json
{
  "reason": "Customer changed plan"
}
```
Response 200:
```json
{
  "id": "guid",
  "status": "Cancelled",
  "note": "optional",
  "totalAmount": 80000
}
```

POST /api/bookings/{id}/confirm
Auth: Staff/Admin flow
Request body: không có
Response 200:
```json
{
  "id": "guid",
  "status": "Confirmed",
  "totalAmount": 80000
}
```

POST /api/bookings/{id}/start
Auth: Staff/Admin flow
Request body: không có
Response 200:
```json
{
  "id": "guid",
  "status": "InProgress",
  "totalAmount": 80000
}
```

POST /api/bookings/{id}/complete
Auth: Staff/Admin flow
Request body: không có
Response 200:
```json
{
  "id": "guid",
  "status": "Completed",
  "totalAmount": 80000
}
```

13.8 Payments

GET /api/payments/{id}
Auth: None/Bearer
Request body: không có
Response 200:
```json
{
  "id": "guid",
  "bookingId": "guid",
  "amount": 80000,
  "method": "Cash",
  "status": "Pending",
  "createdAt": "2026-07-01T00:00:00Z",
  "paidAt": null,
  "referenceNumber": null,
  "note": "Pay at counter"
}
```

POST /api/payments
Auth: None/Bearer
Request:
```json
{
  "bookingId": "guid",
  "amount": 80000,
  "method": "Cash",
  "note": "Pay at counter"
}
```
Response 201:
```json
{
  "id": "guid",
  "bookingId": "guid",
  "amount": 80000,
  "method": "Cash",
  "status": "Pending",
  "createdAt": "2026-07-01T00:00:00Z",
  "paidAt": null,
  "referenceNumber": null,
  "note": "Pay at counter"
}
```

POST /api/payments/{id}/paid
Auth: Staff/Admin flow
Request:
```json
{
  "referenceNumber": "REF001",
  "note": "Paid at counter"
}
```
Response 200:
```json
{
  "id": "guid",
  "bookingId": "guid",
  "amount": 80000,
  "method": "Cash",
  "status": "Paid",
  "createdAt": "2026-07-01T00:00:00Z",
  "paidAt": "2026-07-01T09:30:00Z",
  "referenceNumber": "REF001",
  "note": "Pay at counter; Paid at counter"
}
```

POST /api/payments/{id}/void
Auth: Staff/Admin flow
Request:
```json
{
  "note": "Customer cancelled before payment"
}
```
Response 200:
```json
{
  "id": "guid",
  "bookingId": "guid",
  "amount": 80000,
  "method": "Cash",
  "status": "Voided",
  "createdAt": "2026-07-01T00:00:00Z",
  "paidAt": null,
  "referenceNumber": null,
  "note": "Customer cancelled before payment"
}
```

13.9 Loyalty - Settings, tiers, points

GET /api/loyalty/settings
Auth: None/Bearer
Response 200:
```json
{
  "pointEarnRateAmount": 10000,
  "pointEarnRatePoints": 1,
  "pointExpiryMonths": 12,
  "earnRule": "1 point per 10,000 VND, multiplied by current tier."
}
```

GET /api/loyalty/tiers?page=1&pageSize=20&includeInactive=false
Auth: None/Bearer
Response 200:
```json
{
  "items": [
    {
      "id": "guid",
      "name": "Silver",
      "rank": 2,
      "minSpent": 1000000,
      "minVisits": 5,
      "qualificationPeriodMonths": 12,
      "qualificationMode": "AllConditions",
      "bookingWindowDays": 7,
      "priorityLevel": 1,
      "pointMultiplier": 1.2,
      "benefits": "5% discount",
      "status": "Active"
    }
  ],
  "page": 1,
  "pageSize": 20,
  "totalCount": 1
}
```

GET /api/loyalty/tiers/{id}
Auth: None/Bearer
Response 200:
```json
{
  "id": "guid",
  "name": "Silver",
  "rank": 2,
  "minSpent": 1000000,
  "minVisits": 5,
  "qualificationPeriodMonths": 12,
  "qualificationMode": "AllConditions",
  "bookingWindowDays": 7,
  "priorityLevel": 1,
  "pointMultiplier": 1.2,
  "benefits": "5% discount",
  "status": "Active"
}
```

POST /api/loyalty/tiers
Auth: Admin flow
Request:
```json
{
  "name": "Silver",
  "rank": 2,
  "minSpent": 1000000,
  "minVisits": 5,
  "qualificationPeriodMonths": 12,
  "qualificationMode": "AllConditions",
  "bookingWindowDays": 7,
  "priorityLevel": 1,
  "pointMultiplier": 1.2,
  "benefits": "5% discount",
  "isActive": true
}
```
Response 201: `LoyaltyTierResponse`

PUT /api/loyalty/tiers/{id}
Auth: Admin flow
Request: giống POST /api/loyalty/tiers
Response 200: `LoyaltyTierResponse`

DELETE /api/loyalty/tiers/{id}
Auth: Admin flow
Response 200:
```json
true
```

GET /api/loyalty/customers/{customerId}/points/balance
Auth: None/Bearer
Response 200:
```json
{
  "customerId": "guid",
  "currentPoints": 100,
  "lifetimePoints": 200,
  "currentTier": "Silver",
  "totalSpent": 500000,
  "totalVisits": 5
}
```

GET /api/loyalty/customers/{customerId}/points/history?page=1&pageSize=20
Auth: None/Bearer
Response 200:
```json
{
  "items": [
    {
      "id": "guid",
      "customerId": "guid",
      "bookingId": "guid",
      "washHistoryId": "guid",
      "redemptionId": null,
      "points": 8,
      "originalPoints": 8,
      "remainingPoints": 8,
      "balanceAfter": 108,
      "type": "Earn",
      "expiryDate": "2027-07-01T00:00:00Z",
      "idempotencyKey": "wash:guid:earn",
      "description": "Points earned from completed wash",
      "createdAt": "2026-07-01T00:00:00Z"
    }
  ],
  "page": 1,
  "pageSize": 20,
  "totalCount": 1
}
```

13.10 Loyalty - Rewards

GET /api/loyalty/rewards?page=1&pageSize=20&includeInactive=false
Auth: None/Bearer
Response 200:
```json
{
  "items": [
    {
      "id": "guid",
      "name": "Voucher 50k",
      "description": "Discount",
      "type": "FixedDiscount",
      "pointsRequired": 100,
      "value": 50000,
      "serviceId": null,
      "validFrom": "2026-07-01T00:00:00Z",
      "validTo": "2026-12-31T00:00:00Z",
      "usageLimitPerCustomer": 1,
      "status": "Active",
      "createdAt": "2026-07-01T00:00:00Z"
    }
  ],
  "page": 1,
  "pageSize": 20,
  "totalCount": 1
}
```

GET /api/loyalty/rewards/{id}
Auth: None/Bearer
Response 200: `RewardResponse`

POST /api/loyalty/rewards
Auth: Admin flow
Request:
```json
{
  "name": "Voucher 50k",
  "description": "Discount",
  "type": "FixedDiscount",
  "pointsRequired": 100,
  "value": 50000,
  "serviceId": null,
  "validFrom": "2026-07-01T00:00:00Z",
  "validTo": "2026-12-31T00:00:00Z",
  "usageLimitPerCustomer": 1,
  "isActive": true
}
```
Response 201: `RewardResponse`

PUT /api/loyalty/rewards/{id}
Auth: Admin flow
Request: giống POST /api/loyalty/rewards
Response 200: `RewardResponse`

DELETE /api/loyalty/rewards/{id}
Auth: Admin flow
Response 200:
```json
true
```

POST /api/loyalty/rewards/{id}/redeem
Auth: Customer/Admin flow
Request:
```json
{
  "customerId": "guid",
  "bookingId": "guid",
  "idempotencyKey": "client-generated-key"
}
```
Response 201:
```json
{
  "id": "guid",
  "customerId": "guid",
  "rewardId": "guid",
  "bookingId": "guid",
  "pointsSpent": 100,
  "status": "Reserved",
  "redeemedAt": "2026-07-01T00:00:00Z",
  "expiresAt": "2026-12-31T00:00:00Z",
  "usedAt": null
}
```

13.11 Loyalty - Promotions

GET /api/loyalty/promotions?page=1&pageSize=20&includeInactive=false
Auth: None/Bearer
Response 200:
```json
{
  "items": [
    {
      "id": "guid",
      "name": "Summer Sale",
      "code": "SUMMER10",
      "description": "10% off",
      "type": "PercentageDiscount",
      "value": 10,
      "maxDiscountAmount": 50000,
      "bonusPoints": 0,
      "freeServiceId": null,
      "minimumSpend": 100000,
      "startDate": "2026-07-01T00:00:00Z",
      "endDate": "2026-07-31T23:59:59Z",
      "minTierId": null,
      "totalUsageLimit": 100,
      "usageLimitPerCustomer": 1,
      "priority": 1,
      "isStackable": false,
      "status": "Active",
      "createdAt": "2026-07-01T00:00:00Z",
      "serviceIds": ["guid"]
    }
  ],
  "page": 1,
  "pageSize": 20,
  "totalCount": 1
}
```

GET /api/loyalty/promotions/{id}
Auth: None/Bearer
Response 200: `PromotionResponse`

POST /api/loyalty/promotions
Auth: Admin flow
Request:
```json
{
  "name": "Summer Sale",
  "code": "SUMMER10",
  "description": "10% off",
  "type": "PercentageDiscount",
  "value": 10,
  "maxDiscountAmount": 50000,
  "bonusPoints": 0,
  "freeServiceId": null,
  "minimumSpend": 100000,
  "startDate": "2026-07-01T00:00:00Z",
  "endDate": "2026-07-31T23:59:59Z",
  "minTierId": null,
  "totalUsageLimit": 100,
  "usageLimitPerCustomer": 1,
  "priority": 1,
  "isStackable": false,
  "isActive": true,
  "serviceIds": ["guid"]
}
```
Response 201: `PromotionResponse`

PUT /api/loyalty/promotions/{id}
Auth: Admin flow
Request: giống POST /api/loyalty/promotions
Response 200: `PromotionResponse`

DELETE /api/loyalty/promotions/{id}
Auth: Admin flow
Response 200:
```json
true
```

POST /api/loyalty/promotions/{id}/send
Auth: Admin flow
Request:
```json
{
  "customerIds": ["guid"],
  "expiresAt": "2026-07-31T23:59:59Z"
}
```
Response 200:
```json
{
  "promotionId": "guid",
  "sentCount": 1,
  "skippedCount": 0
}
```

POST /api/loyalty/promotions/{id}/apply
Auth: Customer/Admin flow
Request:
```json
{
  "bookingId": "guid",
  "customerId": "guid",
  "code": "SUMMER10"
}
```
Response 200:
```json
{
  "bookingId": "guid",
  "promotionId": "guid",
  "discountAmount": 50000,
  "bonusPoints": 0,
  "totalBeforeDiscount": 150000,
  "totalAfterDiscount": 100000
}
```

13.12 Loyalty - Tier history, wash history, dashboard

GET /api/loyalty/wash-history?customerId={guid}&page=1&pageSize=20
Auth: None/Bearer
Response 200:
```json
{
  "items": [
    {
      "id": "guid",
      "bookingId": "guid",
      "washDate": "2026-07-01T09:30:00Z",
      "actualTotalAmount": 150000,
      "discountAmount": 50000,
      "finalAmount": 100000,
      "pointsEarned": 10,
      "rewardUsed": 50000,
      "customerRating": 5,
      "feedback": "Good",
      "createdAt": "2026-07-01T09:30:00Z"
    }
  ],
  "page": 1,
  "pageSize": 20,
  "totalCount": 1
}
```

POST /api/loyalty/customers/{customerId}/tier/evaluate
Auth: Admin flow
Request body: không có
Response 200:
```json
{
  "customerId": "guid",
  "previousTierId": "guid",
  "currentTierId": "guid",
  "currentTierName": "Silver",
  "qualifiedSpent": 1000000,
  "qualifiedVisits": 5,
  "changed": true
}
```

GET /api/loyalty/customers/{customerId}/tier/history?page=1&pageSize=20
Auth: None/Bearer
Response 200:
```json
{
  "items": [
    {
      "id": "guid",
      "customerId": "guid",
      "previousTierId": "guid",
      "newTierId": "guid",
      "reviewPeriodStart": "2026-06-01T00:00:00Z",
      "reviewPeriodEnd": "2026-06-30T23:59:59Z",
      "qualifiedSpent": 1000000,
      "qualifiedVisits": 5,
      "changeReason": "MonthlyReview",
      "changedAt": "2026-07-01T00:00:00Z",
      "notes": null
    }
  ],
  "page": 1,
  "pageSize": 20,
  "totalCount": 1
}
```

GET /api/loyalty/dashboard?fromDate=2026-07-01&toDate=2026-07-31
Auth: Admin flow
Response 200:
```json
{
  "activeCustomers": 10,
  "activeRewards": 3,
  "pointsIssued": 1000,
  "pointsRedeemed": 200,
  "revenue": 5000000,
  "completedWashes": 50
}
```

13.13 Wash histories

GET /api/wash-histories/me?page=1&pageSize=10
Auth: CustomerOnly
Response 200:
```json
{
  "success": true,
  "data": {
    "items": [
      {
        "washHistoryID": "guid",
        "bookingID": "guid",
        "washDate": "2026-07-01T09:30:00Z",
        "finalAmount": 100000,
        "pointsEarned": 10,
        "customerRating": 5,
        "vehiclePlate": "51A12345",
        "branchName": "District 1",
        "services": ["Basic Wash"]
      }
    ],
    "page": 1,
    "pageSize": 10,
    "totalCount": 1
  },
  "message": null
}
```

GET /api/wash-histories/me/{washHistoryId}
Auth: CustomerOnly
Response 200:
```json
{
  "success": true,
  "data": {
    "washHistoryID": "guid",
    "bookingID": "guid",
    "washDate": "2026-07-01T09:30:00Z",
    "actualTotalAmount": 150000,
    "discountAmount": 50000,
    "finalAmount": 100000,
    "pointsEarned": 10,
    "rewardUsed": 50000,
    "customerRating": 5,
    "feedback": "Good",
    "vehiclePlate": "51A12345",
    "branchName": "District 1",
    "services": [
      {
        "serviceID": "guid",
        "serviceName": "Basic Wash",
        "price": 100000
      }
    ]
  },
  "message": null
}
```

GET /api/wash-histories/customer/{customerId}?page=1&pageSize=10
Auth: StaffOrAdmin
Response 200: giống list item của `/api/wash-histories/me`

13.14 AI

POST /api/ai/chat
Auth: CustomerOnly + rate limit
Request:
```json
{
  "message": "Tôi nên rửa xe gói nào?",
  "conversationId": null
}
```
Response 200:
```json
{
  "reply": "Bạn có thể chọn Basic Wash nếu muốn nhanh và tiết kiệm.",
  "conversationId": "conversation-id",
  "isFallback": false,
  "source": "gemini"
}
```

POST /api/ai/suggest-services
Auth: CustomerOnly + rate limit
Request:
```json
{
  "vehicleType": "Sedan",
  "preference": "cheap and fast"
}
```
Response 200:
```json
{
  "suggestions": [
    {
      "serviceId": "guid",
      "serviceName": "Basic Wash",
      "price": 80000,
      "reason": "Phù hợp nhu cầu nhanh và tiết kiệm."
    }
  ],
  "summary": "Gợi ý 1 dịch vụ phù hợp.",
  "isFallback": false,
  "source": "gemini"
}
```

POST /api/ai/admin/chat
Auth: AdminOnly + rate limit
Request:
```json
{
  "message": "Tóm tắt doanh thu tháng này",
  "conversationId": null
}
```
Response 200:
```json
{
  "reply": "Doanh thu tháng này là ...",
  "conversationId": "conversation-id",
  "isFallback": false,
  "source": "gemini"
}
```

13.15 Admin users và behavioral logs

PUT /api/admin/users/{userId}/status
Auth: AdminOnly
Request:
```json
{
  "status": "Active"
}
```
Response 200:
```json
{
  "userID": "guid",
  "username": "customer01",
  "fullName": "Nguyen Van A",
  "email": "a@example.com",
  "role": "Customer",
  "status": "Active",
  "emailVerified": true,
  "createdAt": "2026-07-01T00:00:00Z"
}
```

GET /api/admin/behavioral-logs?customerID={guid}&actionType=Book&from=2026-07-01&to=2026-07-31&page=1&pageSize=20
Auth: AdminOnly
Request body: không có
Response 200:
```json
{
  "success": true,
  "data": {
    "items": [
      {
        "logID": "guid",
        "customerID": "guid",
        "customerName": "Nguyen Van A",
        "actionType": "Book",
        "actionTime": "2026-07-01T09:00:00Z",
        "pointsChanged": 10,
        "spendingAmount": 100000,
        "notes": "Booking created"
      }
    ],
    "page": 1,
    "pageSize": 20,
    "totalCount": 1,
    "totalPages": 1,
    "hasNext": false,
    "hasPrevious": false
  },
  "message": null
}
```

GET /api/admin/behavioral-logs/export?customerID={guid}&actionType=Book&from=2026-07-01&to=2026-07-31
Auth: AdminOnly
Response: file CSV `behavioral-logs.csv`
```csv
LogID,CustomerID,CustomerName,ActionType,ActionTime,PointsChanged,SpendingAmount,Notes
guid,guid,Nguyen Van A,Book,2026-07-01T09:00:00Z,10,100000,Booking created
```

13.16 Legacy Services controller

Controller legacy `API/Controllers/ServicesController.cs` có route thực tế `/api/Services`, có thể trùng với `/api/services` của Operations vì route ASP.NET thường không phân biệt hoa/thường.

FE nên ưu tiên `/api/services` của Operations. Nếu backend vẫn giữ controller legacy, các endpoint legacy gồm:
```http
GET /api/Services
GET /api/Services/{id}
POST /api/Services
PUT /api/Services/{id}
PATCH /api/Services/{id}
DELETE /api/Services/{id}
```
"""


def add_code_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    for idx, line in enumerate(text.splitlines()):
        if idx:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.5)


def add_text_paragraph(document: Document, text: str) -> None:
    if not text.strip():
        document.add_paragraph()
        return
    if text.startswith("13."):
        level = 1 if text == "13. Phụ lục JSON chi tiết cho từng API" else 2
        paragraph = document.add_heading(text, level=level)
        paragraph.paragraph_format.keep_with_next = True
        return
    if text.startswith(("GET ", "POST ", "PUT ", "DELETE ", "PATCH ")):
        paragraph = document.add_heading(text, level=3)
        paragraph.paragraph_format.keep_with_next = True
        return
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(5)


def append_to_docx() -> None:
    document = Document(DOCX_PATH)
    document.add_page_break()
    in_code = False
    code_lines: list[str] = []
    for raw_line in APPENDIX.strip().splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            if in_code:
                add_code_paragraph(document, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
        else:
            add_text_paragraph(document, line)
    if code_lines:
        add_code_paragraph(document, "\n".join(code_lines))
    document.save(DOCX_PATH)


def main() -> None:
    APPENDIX_TXT_PATH.parent.mkdir(parents=True, exist_ok=True)
    APPENDIX_TXT_PATH.write_text("\n\n" + APPENDIX.strip() + "\n", encoding="utf-8")
    append_to_docx()
    print(f"updated={DOCX_PATH}")
    print(f"appendix_text={APPENDIX_TXT_PATH}")


if __name__ == "__main__":
    main()
