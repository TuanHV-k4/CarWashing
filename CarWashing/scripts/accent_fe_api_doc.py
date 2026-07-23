from __future__ import annotations

from pathlib import Path

from docx import Document


DOCX_PATH = Path("AutoWash_Pro_API_FE_Chi_Tiet.docx")

REPLACEMENTS = {
    "AutoWash Pro - API Contract va Luong Tich Hop FE": "AutoWash Pro - API Contract và Luồng Tích Hợp FE",
    "Tai lieu danh cho Frontend: endpoint, payload, response, auth, enum, luong man hinh va cac luu y tich hop.": "Tài liệu dành cho Frontend: endpoint, payload, response, auth, enum, luồng màn hình và các lưu ý tích hợp.",
    "1. Quy uoc chung": "1. Quy ước chung",
    "Request/response mac dinh la JSON. Header chung: Content-Type: application/json.": "Request/response mặc định là JSON. Header chung: Content-Type: application/json.",
    "Endpoint co auth can them Authorization: Bearer <accessToken>.": "Endpoint có auth cần thêm Authorization: Bearer <accessToken>.",
    "Enum serialize theo ten chuoi do backend dung JsonStringEnumConverter; FE gui chuoi enum, khong gui so.": "Enum serialize theo tên chuỗi do backend dùng JsonStringEnumConverter; FE gửi chuỗi enum, không gửi số.",
    "DateTime gui theo ISO 8601, uu tien UTC, vi du 2026-07-01T09:00:00Z. TimeSpan cua branch gui dang 08:00:00.": "DateTime gửi theo ISO 8601, ưu tiên UTC, ví dụ 2026-07-01T09:00:00Z. TimeSpan của branch gửi dạng 08:00:00.",
    "Loi validation/business cua Operations/Loyalty thuong tra application/problem+json voi title/detail/status.": "Lỗi validation/business của Operations/Loyalty thường trả application/problem+json với title/detail/status.",
    "Operations va Loyalty hien khong gan [Authorize] tren controller, nhung mot so ham booking van doc current customer tu JWT neu co. FE nen van gui token sau login cho cac luong customer.": "Operations và Loyalty hiện không gắn [Authorize] trên controller, nhưng một số hàm booking vẫn đọc current customer từ JWT nếu có. FE nên vẫn gửi token sau login cho các luồng customer.",
    "ProblemDetails mau:": "ProblemDetails mẫu:",
    "2. Auth va phan quyen": "2. Auth và phân quyền",
    "Auth DTO chinh": "Auth DTO chính",
    "Luong FE de xac thuc": "Luồng FE để xác thực",
    "Man hinh Register goi POST /api/Auth/register; neu thanh cong chuyen sang man hinh OTP.": "Màn hình Register gọi POST /api/Auth/register; nếu thành công chuyển sang màn hình OTP.",
    "Man hinh OTP goi POST /api/Auth/verify-email. Neu OTP sai/hết hạn hien message detail; co nut goi resend-otp.": "Màn hình OTP gọi POST /api/Auth/verify-email. Nếu OTP sai/hết hạn hiển thị message detail; có nút gọi resend-otp.",
    "Man hinh Login goi POST /api/Auth/login, luu accessToken va role. Moi request can auth gan Authorization header.": "Màn hình Login gọi POST /api/Auth/login, lưu accessToken và role. Mọi request cần auth gắn Authorization header.",
    "Khi reload app, FE co the goi GET /api/Auth/me de verify token va khoi phuc session.": "Khi reload app, FE có thể gọi GET /api/Auth/me để verify token và khôi phục session.",
    "3. Customer va Vehicle": "3. Customer và Vehicle",
    "4. Operations: catalog, chi nhanh, bay rua": "4. Operations: catalog, chi nhánh, bay rửa",
    "Luong dat lich FE": "Luồng đặt lịch FE",
    "Load profile/customer va danh sach xe: GET /api/customers/me, GET /api/vehicles/me.": "Load profile/customer và danh sách xe: GET /api/customers/me, GET /api/vehicles/me.",
    "Load catalog: GET /api/services va GET /api/branches; khi chon branch co the load GET /api/wash-bays de hien tinh trang noi bo neu can.": "Load catalog: GET /api/services và GET /api/branches; khi chọn branch có thể load GET /api/wash-bays để hiển thị tình trạng nội bộ nếu cần.",
    "FE tinh preview end time = bookingStartTime + durationMinutes cua service, nhung backend moi la nguon chinh xac.": "FE tính preview end time = bookingStartTime + durationMinutes của service, nhưng backend mới là nguồn chính xác.",
    "Goi POST /api/bookings. Neu 409 'No available wash bay...' hien thong bao chon khung gio khac.": "Gọi POST /api/bookings. Nếu 409 'No available wash bay...' hiển thị thông báo chọn khung giờ khác.",
    "Sau khi tao booking, co the hien trang thanh toan hoac trang chi tiet booking.": "Sau khi tạo booking, có thể hiển thị trang thanh toán hoặc trang chi tiết booking.",
    "Luong trang thai booking": "Luồng trạng thái booking",
    "Pending/Confirmed/InProgress -> Cancelled neu chua paid va chua completed": "Pending/Confirmed/InProgress -> Cancelled nếu chưa paid và chưa completed",
    "7. Loyalty, rewards va promotions": "7. Loyalty, rewards và promotions",
    "Loyalty rules can biet": "Loyalty rules cần biết",
    "Luong loyalty khi hoan tat rua xe": "Luồng loyalty khi hoàn tất rửa xe",
    "Staff/Admin confirm va start booking theo workflow operations.": "Staff/Admin confirm và start booking theo workflow operations.",
    "Khi complete booking, backend tao WashHistory, tinh discount/bonus tu BookingPromotions, cong diem cho customer.": "Khi complete booking, backend tạo WashHistory, tính discount/bonus từ BookingPromotions, cộng điểm cho customer.",
    "Backend ghi LoyaltyPointTransaction Earn voi idempotencyKey wash:{bookingId}:earn de tranh cong lap.": "Backend ghi LoyaltyPointTransaction Earn với idempotencyKey wash:{bookingId}:earn để tránh cộng lặp.",
    "Backend cap nhat TotalSpent, TotalVisits, LastVisitDate va evaluate tier.": "Backend cập nhật TotalSpent, TotalVisits, LastVisitDate và evaluate tier.",
    "FE refresh GET /api/customers/me, points/balance, wash-history de cap nhat UI.": "FE refresh GET /api/customers/me, points/balance, wash-history để cập nhật UI.",
    "10. Admin users va behavioral logs": "10. Admin users và behavioral logs",
    "12. Phu luc enum FE nen map": "12. Phụ lục enum FE nên map",
    "Pham vi va nguon doi chieu": "Phạm vi và nguồn đối chiếu",
    "Tai lieu duoc tong hop tu controller, DTO, enum va service logic trong repo hien tai. Base URL local thuong la https://localhost:<port> hoac http://localhost:<port>. Swagger co tai /swagger khi API chay.": "Tài liệu được tổng hợp từ controller, DTO, enum và service logic trong repo hiện tại. Base URL local thường là https://localhost:<port> hoặc http://localhost:<port>. Swagger có tại /swagger khi API chạy.",
    "Dang ky customer, backend tao user/customer va gui OTP email.": "Đăng ký customer, backend tạo user/customer và gửi OTP email.",
    "Xac thuc email bang OTP.": "Xác thực email bằng OTP.",
    "Gui lai OTP khi het han/khong nhan duoc.": "Gửi lại OTP khi hết hạn/không nhận được.",
    "Dang nhap, nhan JWT de FE luu va gan vao header.": "Đăng nhập, nhận JWT để FE lưu và gắn vào header.",
    "Kiem tra token va lay claims user hien tai.": "Kiểm tra token và lấy claims user hiện tại.",
    "Lay profile, diem, tier hien tai va tier perks.": "Lấy profile, điểm, tier hiện tại và tier perks.",
    "Danh sach xe cua customer hien tai.": "Danh sách xe của customer hiện tại.",
    "Them xe.": "Thêm xe.",
    "Sua thong tin xe.": "Sửa thông tin xe.",
    "Doi status xe.": "Đổi status xe.",
    "Bien so, FE nen normalize upper-case neu can.": "Biển số, FE nên normalize upper-case nếu cần.",
    "Thong tin mo ta xe.": "Thông tin mô tả xe.",
    "Lay goi dich vu dang ban.": "Lấy gói dịch vụ đang bán.",
    "Chi tiet dich vu.": "Chi tiết dịch vụ.",
    "Tao dich vu.": "Tạo dịch vụ.",
    "Sua dich vu.": "Sửa dịch vụ.",
    "Inactive/archive dich vu.": "Inactive/archive dịch vụ.",
    "Lay chi nhanh dang mo.": "Lấy chi nhánh đang mở.",
    "Lay bay rua theo chi nhanh.": "Lấy bay rửa theo chi nhánh.",
    "Canh bao route trung api/services": "Cảnh báo route trùng api/services",
    "Repo hien co them controller legacy API/Controllers/ServicesController.cs voi route api/[controller] => /api/Services. ASP.NET route thuong khong phan biet hoa/thuong, co the trung voi /api/services cua Operations. FE nen uu tien Operations Services va backend nen doi/xoa route legacy de tranh ambiguous endpoint.": "Repo hiện có thêm controller legacy API/Controllers/ServicesController.cs với route api/[controller] => /api/Services. ASP.NET route thường không phân biệt hoa/thường, có thể trùng với /api/services của Operations. FE nên ưu tiên Operations Services và backend nên đổi/xóa route legacy để tránh ambiguous endpoint.",
    "Danh sach booking; customer chi thay cua minh neu token co customerId, Admin/Staff thay rong hon.": "Danh sách booking; customer chỉ thấy của mình nếu token có customerId, Admin/Staff thấy rộng hơn.",
    "Chi tiet booking kem branchName/washBayName.": "Chi tiết booking kèm branchName/washBayName.",
    "Bearer nen co": "Bearer nên có",
    "Customer tao lich rua xe.": "Customer tạo lịch rửa xe.",
    "Huy booking khi chua completed va chua paid.": "Hủy booking khi chưa completed và chưa paid.",
    "InProgress -> Completed, tao wash history va cong diem.": "InProgress -> Completed, tạo wash history và cộng điểm.",
    "Phai thuoc customer hien tai, lay tu GET /api/vehicles/me.": "Phải thuộc customer hiện tại, lấy từ GET /api/vehicles/me.",
    "Chi nhanh phai Open, lay tu GET /api/branches.": "Chi nhánh phải Open, lấy từ GET /api/branches.",
    "Service phai Active, lay tu GET /api/services.": "Service phải Active, lấy từ GET /api/services.",
    "Phai o tuong lai, nam trong gio mo cua branch, khong vuot bookingWindowDays theo tier.": "Phải ở tương lai, nằm trong giờ mở cửa branch, không vượt bookingWindowDays theo tier.",
    "Backend tu tim bay rua kha dung; FE khong can chon neu luong hien tai khong yeu cau.": "Backend tự tìm bay rửa khả dụng; FE không cần chọn nếu luồng hiện tại không yêu cầu.",
    "Lay chi tiet payment.": "Lấy chi tiết payment.",
    "Tao payment pending cho booking.": "Tạo payment pending cho booking.",
    "Danh dau da thanh toan.": "Đánh dấu đã thanh toán.",
    "Huy payment pending.": "Hủy payment pending.",
    "Code error message cu co the ghi Card nhung enum thuc te khong co Card.": "Code error message cũ có thể ghi Card nhưng enum thực tế không có Card.",
    "Phai > 0 va bang booking.EstimatedTotalAmount.": "Phải > 0 và bằng booking.EstimatedTotalAmount.",
    "Chi Pending moi mark paid; Voided khong paid lai duoc.": "Chỉ Pending mới mark paid; Voided không paid lại được.",
    "Chi Pending moi void; Paid khong void duoc.": "Chỉ Pending mới void; Paid không void được.",
    "Lay rule tinh diem.": "Lấy rule tính điểm.",
    "Danh sach tier.": "Danh sách tier.",
    "Quan tri tier.": "Quản trị tier.",
    "So du diem customer.": "Số dư điểm customer.",
    "Lich su diem.": "Lịch sử điểm.",
    "Danh sach reward doi diem.": "Danh sách reward đổi điểm.",
    "Doi diem lay reward.": "Đổi điểm lấy reward.",
    "Danh gia lai tier.": "Đánh giá lại tier.",
    "Chi so loyalty.": "Chỉ số loyalty.",
    "Gan promotion cho danh sach customer.": "Gắn promotion cho danh sách customer.",
    "Ap promotion vao booking pending.": "Áp promotion vào booking pending.",
    "1 point / 10,000 VND * pointMultiplier cua tier + bonusPoints tu promotion.": "1 point / 10,000 VND * pointMultiplier của tier + bonusPoints từ promotion.",
    "Diem earn het han sau 12 thang.": "Điểm earn hết hạn sau 12 tháng.",
    "Tieu diem theo cac transaction Earn con remainingPoints, sap xep theo expiry som nhat.": "Tiêu điểm theo các transaction Earn còn remainingPoints, sắp xếp theo expiry sớm nhất.",
    "AllConditions hoac AnyCondition.": "AllConditions hoặc AnyCondition.",
    "Lich su rua xe cua customer dang dang nhap.": "Lịch sử rửa xe của customer đang đăng nhập.",
    "Chi tiet mot wash history cua customer.": "Chi tiết một wash history của customer.",
    "Staff/Admin xem lich su theo customer.": "Staff/Admin xem lịch sử theo customer.",
    "History trong module loyalty; neu khong truyen customerId tra tat ca.": "History trong module loyalty; nếu không truyền customerId trả tất cả.",
    "Ten field khac nhau giua controller wash-histories va loyalty/wash-history.": "Tên field khác nhau giữa controller wash-histories và loyalty/wash-history.",
    "Tong truoc discount trong DTO loyalty.": "Tổng trước discount trong DTO loyalty.",
    "Tong discount tu promotion/reward.": "Tổng discount từ promotion/reward.",
    "So tien tinh diem va hien thanh toan cuoi.": "Số tiền tính điểm và hiển thị thanh toán cuối.",
    "Diem nhan duoc khi complete.": "Điểm nhận được khi complete.",
    "Chat tu van cho customer.": "Chat tư vấn cho customer.",
    "Goi y goi dich vu theo xe/nhu cau.": "Gợi ý gói dịch vụ theo xe/nhu cầu.",
    "Doi trang thai user.": "Đổi trạng thái user.",
    "Xem log hanh vi.": "Xem log hành vi.",
    "Tai CSV behavioral-logs.csv.": "Tải CSV behavioral-logs.csv.",
    "Dang ky, OTP, login, me.": "Đăng ký, OTP, login, me.",
    "Profile/tier/points tong quan.": "Profile/tier/points tổng quan.",
    "Dat lich va van hanh trang thai.": "Đặt lịch và vận hành trạng thái.",
    "Thu tien tai quay/void.": "Thu tiền tại quầy/void.",
    "Quan tri user/log.": "Quản trị user/log.",
    "Tier, diem, reward, promotion.": "Tier, điểm, reward, promotion.",
    "Lich su rua xe.": "Lịch sử rửa xe.",
    "Chat/goi y.": "Chat/gợi ý.",
}


def replace_in_paragraph(paragraph) -> int:
    changed = 0
    for run in paragraph.runs:
        current = run.text
        updated = current
        for old, new in REPLACEMENTS.items():
            updated = updated.replace(old, new)
        if updated != current:
            run.text = updated
            changed += 1
    return changed


def iter_all_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        for part in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            yield from part.paragraphs
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def main() -> None:
    document = Document(DOCX_PATH)
    changed = sum(replace_in_paragraph(paragraph) for paragraph in iter_all_paragraphs(document))
    document.save(DOCX_PATH)
    print(f"changed_runs={changed}")


if __name__ == "__main__":
    main()
